import os
import pandas as pd
import PyPDF2
from PIL import Image
import pytesseract
from docx import Document
from typing import Dict, Any

def process_csv(file_path: str) -> Dict[str, Any]:
    """Process CSV file"""
    df = pd.read_csv(file_path)
    
    return {
        'type': 'csv',
        'rows': len(df),
        'columns': list(df.columns),
        'preview': df.head(5).to_dict(orient='records'),
        'data': df.to_dict(orient='records'),
        'summary': df.describe().to_dict()
    }

def process_excel(file_path: str) -> Dict[str, Any]:
    """Process Excel file"""
    df = pd.read_excel(file_path)
    
    return {
        'type': 'excel',
        'rows': len(df),
        'columns': list(df.columns),
        'preview': df.head(5).to_dict(orient='records'),
        'data': df.to_dict(orient='records'),
        'summary': df.describe().to_dict()
    }

def process_pdf(file_path: str) -> Dict[str, Any]:
    """Process PDF file"""
    text = ""
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    
    return {
        'type': 'pdf',
        'pages': num_pages,
        'text': text,
        'preview': text[:500] + '...' if len(text) > 500 else text
    }

def process_image(file_path: str, question: str = None) -> Dict[str, Any]:
    """Process image with OCR only (vision disabled due to quota)"""
    try:
        from PIL import Image
        import pytesseract
        
        image = Image.open(file_path)
        
        # OCR for text extraction
        text = pytesseract.image_to_string(image)
        
        result = {
            'type': 'image',
            'size': image.size,
            'format': image.format,
            'ocr_text': text,
            'preview': text[:500] + '...' if len(text) > 500 else text,
            'note': 'Vision analysis temporarily disabled (Bedrock quota). OCR text extraction active.'
        }
        
        return result
        
    except Exception as e:
        return {
            'type': 'image',
            'error': str(e),
            'text': ''
        }

def process_docx(file_path: str) -> Dict[str, Any]:
    """Process Word document"""
    doc = Document(file_path)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    return {
        'type': 'docx',
        'paragraphs': len(doc.paragraphs),
        'text': text,
        'preview': text[:500] + '...' if len(text) > 500 else text
    }

def process_txt(file_path: str) -> Dict[str, Any]:
    """Process text file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    
    return {
        'type': 'txt',
        'text': text,
        'preview': text[:500] + '...' if len(text) > 500 else text
    }

def process_file(file_path: str, question: str = None) -> Dict[str, Any]:
    """Route to appropriate processor based on file extension"""
    ext = os.path.splitext(file_path)[1].lower()
    
    processors = {
        '.csv': process_csv,
        '.xlsx': process_excel,
        '.xls': process_excel,
        '.pdf': process_pdf,
        '.png': lambda fp: process_image(fp, question),  # ← Pass question
        '.jpg': lambda fp: process_image(fp, question),  # ← Pass question
        '.jpeg': lambda fp: process_image(fp, question), # ← Pass question
        '.docx': process_docx,
        '.txt': process_txt,
    }
    
    processor = processors.get(ext)
    
    if not processor:
        return {'type': 'unknown', 'error': f'Unsupported file type: {ext}'}
    
    try:
        return processor(file_path)
    except Exception as e:
        return {'type': 'error', 'error': str(e)}
    


